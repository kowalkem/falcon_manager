import os
import datetime
from django.conf import settings
from django.http import response
from django.views import generic
from django.views.generic.base import TemplateView
from django.shortcuts import render
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.core.files import File
from docx import Document
from .models import Falcon, Pair, Aviary, Office, Birth_cert, CITES
from .forms import (
    FalconCreateForm,
    FalconUpdateForm,
    PairCreateForm,
    PairUpdateForm,
    YoungFalconCreateForm,
    AviaryCreateForm,
    AviaryUpdateForm,
    OfficeCreateForm,
    OfficeUpdateForm,
    Birth_certCreateForm,
    Birth_certUpdateForm,
    CITESCreateForm,
    CITESUpdateForm,
)

# utility function to remove row from docx table


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def index(request):
    """ Renders index page for breeding app """
    return render(request, "index.html")


class FalconList(LoginRequiredMixin, generic.ListView):
    """ Renders list of all falcons """

    model = Falcon
    paginate_by = 10

    def get_queryset(self):
        return Falcon.objects.filter(owner=self.request.user)

    def get_context_data(self, **kwargs):
        context = super(FalconList, self).get_context_data(**kwargs)
        context["title"] = "List of all falcons"
        return context


class FalconDetail(LoginRequiredMixin, UserPassesTestMixin, generic.DetailView):
    """ Renders info about falcon with specified id """

    model = Falcon

    def test_func(self):
        falcon = self.get_object()
        return self.request.user == falcon.owner

    def get_context_data(self, **kwargs):
        context = super(FalconDetail, self).get_context_data(**kwargs)
        context["title"] = "Falcon details"
        return context


class FalconCreate(LoginRequiredMixin, generic.edit.CreateView):
    """ View for creating new falcon """

    model = Falcon
    form_class = FalconCreateForm

    def form_valid(self, form):
        form.instance.owner = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(FalconCreate, self).get_context_data(**kwargs)
        context["title"] = "Add a new falcon"
        return context


class FalconUpdate(LoginRequiredMixin, UserPassesTestMixin, generic.edit.UpdateView):
    """ View for updating a falcon """

    model = Falcon
    form_class = FalconUpdateForm

    def test_func(self):
        falcon = self.get_object()
        return self.request.user == falcon.owner

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(FalconUpdate, self).get_context_data(**kwargs)
        context["title"] = "Update a falcon"
        return context


class FalconDelete(LoginRequiredMixin, UserPassesTestMixin, generic.DeleteView):
    """ Delete a falcon with a specified id """

    model = Falcon
    success_url = "/breeding/falcons"

    def test_func(self):
        falcon = self.get_object()
        return self.request.user == falcon.owner

    def get_context_data(self, **kwargs):
        context = super(FalconDelete, self).get_context_data(**kwargs)
        context["title"] = "Delete a falcon"
        return context


class YoungFalconCreate(LoginRequiredMixin, generic.edit.CreateView):
    """ View for creating new falcon from pair """

    model = Falcon
    form_class = YoungFalconCreateForm

    def form_valid(self, form):
        form.instance.owner = self.request.user
        pair = Pair.objects.get(pk=self.kwargs["pair_pk"])
        form.instance.father = Falcon.objects.get(pk=pair.male.id)
        form.instance.mother = Falcon.objects.get(pk=pair.female.id)
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(YoungFalconCreate, self).get_context_data(**kwargs)
        context["title"] = "Add a new falcon from specified pair"
        return context


class PairList(LoginRequiredMixin, generic.ListView):
    """ Renders list of all pairs """

    model = Pair
    paginate_by = 10

    def get_queryset(self):
        return Pair.objects.filter(
            male__owner=self.request.user, female__owner=self.request.user
        )

    def get_context_data(self, **kwargs):
        context = super(PairList, self).get_context_data(**kwargs)
        context["title"] = "List of all pairs"
        return context


class PairDetail(LoginRequiredMixin, UserPassesTestMixin, generic.DetailView):
    """ Renders info about pair with specified id """

    model = Pair

    def test_func(self):
        pair = self.get_object()
        return self.request.user == pair.male.owner and pair.female.owner

    def get_context_data(self, **kwargs):
        context = super(PairDetail, self).get_context_data(**kwargs)
        context["title"] = "Pair details"
        return context


class PairCreate(LoginRequiredMixin, generic.edit.CreateView):
    """ View for creating new pair """

    model = Pair
    form_class = PairCreateForm

    def form_valid(self, form):
        form.instance.owner = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(PairCreate, self).get_context_data(**kwargs)
        context["form"].fields["male"].queryset = Falcon.objects.filter(
            sex="M")
        context["form"].fields["female"].queryset = Falcon.objects.filter(
            sex="F")
        context["title"] = "Add a new pair"
        return context


class PairUpdate(LoginRequiredMixin, UserPassesTestMixin, generic.edit.UpdateView):
    """ View for updating a pair """

    model = Pair
    form_class = PairUpdateForm

    def test_func(self):
        pair = self.get_object()
        return self.request.user == pair.male.owner and pair.female.owner

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(PairUpdate, self).get_context_data(**kwargs)
        context["title"] = "Update a pair"
        return context


class PairDelete(LoginRequiredMixin, UserPassesTestMixin, generic.DeleteView):
    """ Delete a pair with a specified id """

    model = Pair
    success_url = "/breeding/pairs"

    def test_func(self):
        pair = self.get_object()
        return self.request.user == pair.male.owner and pair.female.owner

    def get_context_data(self, **kwargs):
        context = super(PairDelete, self).get_context_data(**kwargs)
        context["title"] = "Delete a pair"
        return context


class AviaryList(LoginRequiredMixin, generic.ListView):
    """ Renders list of all aviaries """

    model = Aviary
    paginate_by = 10

    def get_queryset(self):
        return Aviary.objects.filter(owner=self.request.user)

    def get_context_data(self, **kwargs):
        context = super(AviaryList, self).get_context_data(**kwargs)
        context["title"] = "List of all aviaries"
        return context


class AviaryDetail(LoginRequiredMixin, UserPassesTestMixin, generic.DetailView):
    """ Renders info about aviary with specified id """

    model = Aviary

    def test_func(self):
        Aviary = self.get_object()
        return self.request.user == Aviary.owner

    def get_context_data(self, **kwargs):
        context = super(AviaryDetail, self).get_context_data(**kwargs)
        context["title"] = "Aviary details"
        return context


class AviaryCreate(LoginRequiredMixin, generic.edit.CreateView):
    """ View for creating new Aviary """

    model = Aviary
    form_class = AviaryCreateForm

    def form_valid(self, form):
        form.instance.owner = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(AviaryCreate, self).get_context_data(**kwargs)
        context["title"] = "Add a new aviary"
        return context


class AviaryUpdate(LoginRequiredMixin, UserPassesTestMixin, generic.edit.UpdateView):
    """ View for updating a aviary """

    model = Aviary
    form_class = AviaryUpdateForm

    def test_func(self):
        Aviary = self.get_object()
        return self.request.user == Aviary.owner

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(AviaryUpdate, self).get_context_data(**kwargs)
        context["title"] = "Update a Aviary"
        return context


class AviaryDelete(LoginRequiredMixin, UserPassesTestMixin, generic.DeleteView):
    """ Delete an aviary with a specified id """

    model = Aviary
    success_url = "/breeding/aviaries"

    def test_func(self):
        Aviary = self.get_object()
        return self.request.user == Aviary.owner

    def get_context_data(self, **kwargs):
        context = super(AviaryDelete, self).get_context_data(**kwargs)
        context["title"] = "Delete a Aviary"
        return context


class DocList(LoginRequiredMixin, TemplateView):
    template_name = "breeding/doc_list.html"


class DocCreate(LoginRequiredMixin, TemplateView):
    template_name = "breeding/doc_create.html"


class Birth_certList(LoginRequiredMixin, generic.ListView):

    model = Birth_cert
    paginate_by = 10

    def get_queryset(self):
        return Birth_cert.objects.filter(owner=self.request.user)

    def get_context_data(self, **kwargs):
        context = super(Birth_certList, self).get_context_data(**kwargs)
        context["title"] = "List of all birth_certs"
        return context


class Birth_certCreate(LoginRequiredMixin, generic.edit.CreateView):

    model = Birth_cert
    form_class = Birth_certCreateForm

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["request"] = self.request
        return kwargs

    def form_valid(self, form):
        form.instance.owner = self.request.user
        falcons_ids = form.data.getlist('falcons')
        falcons = [Falcon.objects.get(pk=id) for id in falcons_ids]
        res = super().form_valid(form)
        # create docx using form data and append it to form
        doc = Document(os.path.join(settings.MEDIA_ROOT,
                                    'falcon_docs/birth_cert.docx'))
        doc.paragraphs[0].runs[1] = form.instance.document_number

        doc.tables[0].rows[1].cells[1].paragraphs[0].runs[0].text = form.instance.owner.username
        profile = form.instance.owner.profile
        doc.tables[0].rows[2].cells[1].paragraphs[0].runs[
            0].text = f'{profile.street} {profile.house_number}, {profile.zip_code} {profile.city}'

        doc.tables[1].rows[1].cells[1].paragraphs[0].runs[0].text = ''
        doc.tables[1].rows[2].cells[1].paragraphs[0].runs[0].text = ''
        species_list = [falcon.species for falcon in falcons]
        species_count = [(species.name, species.latin, species_list.count(
            species)) for species in set(species_list)]
        for species in species_count:
            doc.tables[1].rows[1].cells[1].paragraphs[0].runs[
                0].text = f'{species[1]} - {species[2]}szt.'
            doc.tables[1].rows[2].cells[1].paragraphs[0].runs[
                0].text = f'{species[0]} - {species[2]}szt.'

        # num females
        doc.tables[3].rows[0].cells[6].paragraphs[0].runs[
            0].text = f'Żeńskich: {len([falcon for falcon in falcons if falcon.sex == "F"])}'
        # num males
        doc.tables[3].rows[1].cells[0].paragraphs[0].runs[
            0].text = f'męskich: {len([falcon for falcon in falcons if falcon.sex == "M"])}'
        # num undefined
        doc.tables[3].rows[1].cells[2].paragraphs[0].runs[
            0].text = f'płci nieznanej: {len([falcon for falcon in falcons if falcon.sex == None])}'
        # num total
        doc.tables[3].rows[1].cells[4].paragraphs[
            0].runs[0].text = f'łącznie: {len(falcons)}'

        # birth dates
        birthdates = [falcon.birth_date for falcon in falcons]
        doc.tables[4].rows[1].cells[0].paragraphs[0].runs[
            0].text = f'urodzenia (wyklucia itp.): {min(birthdates)} - {max(birthdates)}'

        # vet control date
        doc.tables[4].rows[1].cells[1].paragraphs[0].runs[
            0].text = f'kontroli weterynaryjnej miotu: {max(birthdates) + datetime.timedelta(days=5)}'

        # list of falcons
        falcon_list_table_rows = doc.tables[5].rows
        for i, falcon in enumerate(falcons, start=1):
            if len(falcon_list_table_rows) < i+2:
                doc.tables[5].add_row(falcon_list_table_rows[2])
            falcon_list_table_rows[i +
                                   1].cells[0].paragraphs[0].runs[0].text = str(i)
            falcon_list_table_rows[i+1].cells[1].paragraphs[0].runs[0].text = falcon.get_sex_display(
            ) if falcon.sex else 'płeć nieznana'
            falcon_list_table_rows[i +
                                   1].cells[2].paragraphs[0].runs[0].text = f'Obrączka zamknięta nr {falcon.ring}'
            falcon_list_table_rows[i +
                                   1].cells[3].paragraphs[0].runs[0].text = 'brak'

        while len(falcon_list_table_rows) > len(falcons) + 2:
            row_to_delete = falcon_list_table_rows[len(
                falcon_list_table_rows) - 1]
            remove_row(doc.tables[5], row_to_delete)

        # list of parents
        display_data = [[], []]
        pair_list_rows = doc.tables[7].rows

        for youngster_index, falcon in enumerate(falcons, start=1):
            pair = falcon.get_parents()
            if pair in display_data[0]:
                display_data[1][display_data[0].index(
                    pair)].append(youngster_index)
            else:
                display_data[0].append(pair)
                display_data[1].append([youngster_index])
        print(display_data)
        for i, (pair, youngster_indices) in enumerate(zip(display_data[0], display_data[1]), start=1):
            row_index = i*2
            pair_list_rows[row_index].cells[0].paragraphs[0].runs[
                0].text = f'{i}. matka*\n(dla osobników z pozycji {*["6." + str(y) for y in youngster_indices],}'
            pair_list_rows[row_index].cells[1].paragraphs[0].runs[0].text = str(
                pair.female.birth_date)
            pair_list_rows[row_index].cells[2].paragraphs[0].runs[0].text = pair.female.source
            pair_list_rows[row_index].cells[3].paragraphs[0].runs[0].text = pair.female.CITES_num
            pair_list_rows[row_index].cells[3].paragraphs[0].runs[1].text = ""
            row_index = i*2+1
            pair_list_rows[row_index].cells[0].paragraphs[0].runs[
                0].text = f'{i}. ojciec*\n(dla osobników z pozycji {*["6." + str(y) for y in youngster_indices],}'
            pair_list_rows[row_index].cells[1].paragraphs[0].runs[0].text = str(
                pair.male.birth_date)
            pair_list_rows[row_index].cells[2].paragraphs[0].runs[0].text = pair.male.source
            pair_list_rows[row_index].cells[3].paragraphs[0].runs[0].text = pair.male.CITES_num
            pair_list_rows[row_index].cells[3].paragraphs[0].runs[1].text = ""

        while len(pair_list_rows) > len(display_data[0])*2 + 2:
            row_to_delete = pair_list_rows[len(pair_list_rows) - 1]
            remove_row(doc.tables[7], row_to_delete)

        doc.tables[10].rows[0].cells[1].text = f'Powiatowy Lekarz Weterynarii w '

        doc_rel_path = f'falcon_docs/{form.instance.owner.username}_birth_cert_{form.instance.document_number}.docx'
        doc_path = os.path.join(settings.MEDIA_ROOT, doc_rel_path)

        doc.save(doc_path)

        with open(doc_path, 'rb') as f:
            form.instance.cert_file.save(
                f'{form.instance.owner.username}_birth_cert_{form.instance.document_number}.docx', File(f))
            form.instance.save()

        for falcon in falcons:
            falcon.birth_cert = form.instance
            falcon.save()

        return res

    def get_context_data(self, **kwargs):
        context = super(Birth_certCreate, self).get_context_data(**kwargs)
        context["form"].fields["vet_office"].queryset = Office.objects.filter(
            office_type="PIW")
        context["form"].fields["falcons"].queryset = Falcon.objects.all()
        context["title"] = "Add a new birth_cert"
        return context


class Birth_certDetail(LoginRequiredMixin, UserPassesTestMixin, generic.DetailView):

    model = Birth_cert

    def test_func(self):
        Birth_cert = self.get_object()
        return self.request.user == Birth_cert.owner

    def get_context_data(self, **kwargs):
        context = super(Birth_certDetail, self).get_context_data(**kwargs)
        context["title"] = "Birth_cert details"
        return context


class Birth_certUpdate(LoginRequiredMixin, UserPassesTestMixin, generic.edit.UpdateView):

    model = Birth_cert
    form_class = Birth_certUpdateForm

    def test_func(self):
        Birth_cert = self.get_object()
        return self.request.user == Birth_cert.owner

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(Birth_certUpdate, self).get_context_data(**kwargs)
        context["title"] = "Update a Birth_cert"
        return context


class Birth_certDelete(LoginRequiredMixin, UserPassesTestMixin, generic.DeleteView):

    model = Birth_cert
    success_url = "/breeding/birth_certs"

    def test_func(self):
        Birth_cert = self.get_object()
        return self.request.user == Birth_cert.owner

    def get_context_data(self, **kwargs):
        context = super(Birth_certDelete, self).get_context_data(**kwargs)
        context["title"] = "Delete a Birth_cert"
        return context


class CITESList(LoginRequiredMixin, generic.ListView):

    model = CITES
    paginate_by = 10

    def get_queryset(self):
        return CITES.objects.filter(owner=self.request.user)

    def get_context_data(self, **kwargs):
        context = super(CITESList, self).get_context_data(**kwargs)
        context["title"] = "List of all cites"
        return context


class CITESCreate(LoginRequiredMixin, generic.edit.CreateView):

    model = CITES
    form_class = CITESCreateForm

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["request"] = self.request
        return kwargs

    def form_valid(self, form):
        form.instance.owner = self.request.user
        falcons_ids = form.data.getlist('falcons')
        falcons = [Falcon.objects.get(pk=id) for id in falcons_ids]
        res = super().form_valid(form)
        # need to create separate docs for each falcon
        for falcon in falcons:
            # create docx using form data and append it to form
            doc = Document(os.path.join(
                settings.MEDIA_ROOT, 'falcon_docs/cites.docx'))
            # owner data
            doc.tables[0].rows[1].cells[1].paragraphs[1].runs[0].text = form.instance.owner.username
            profile = form.instance.owner.profile
            doc.tables[0].rows[1].cells[1].paragraphs[1].runs[1].text = "\n" + form.instance.owner.profile.street + \
                " " + form.instance.owner.profile.house_number
            doc.tables[0].rows[1].cells[1].paragraphs[1].runs[2].text = "\n" + form.instance.owner.profile.zip_code + \
                " " + form.instance.owner.profile.city
            doc.tables[0].columns[1].cells[29].paragraphs[0].runs[0].text = form.instance.owner.username
            # falcon description
            doc.tables[0].rows[5].cells[1].paragraphs[0].runs[
                3].text = ""
            doc.tables[0].rows[5].cells[1].paragraphs[0].runs[
                4].text = f'\nPłeć - {falcon.sex if falcon.sex else ""}'
            doc.tables[0].rows[5].cells[1].paragraphs[0].runs[
                5].text = f'\nZamknięta obrączka nr: {falcon.ring}'
            doc.tables[0].rows[5].cells[1].paragraphs[0].runs[
                6].text = f'\nData wyklucia: {falcon.birth_date}'
            doc.tables[0].rows[5].cells[1].paragraphs[0].runs[
                7].text = f'\nMatka: {falcon.mother.ring}({falcon.mother.CITES.document_number})'
            doc.tables[0].rows[5].cells[1].paragraphs[0].runs[
                8].text = f'\nOjciec: {falcon.father.ring}({falcon.father.CITES.document_number})'
            # CITES appendix
            doc.tables[0].rows[6].cells[5].paragraphs[0].runs[
                2].text = f'\n{"I" if falcon.species.latin == "Falco peregrinus" else "II"}'
            # species
            doc.tables[0].columns[1].cells[10].paragraphs[
                0].runs[2].text = f'\n{falcon.species.latin}'
            doc.tables[0].columns[1].cells[12].paragraphs[
                0].runs[1].text = f'\n{falcon.species.name}'

            # save doc
            doc_rel_path = f'falcon_docs/{form.instance.owner.username}_CITES_{falcon.ring}.docx'
            doc_path = os.path.join(settings.MEDIA_ROOT, doc_rel_path)

            doc.save(doc_path)

            cites = CITES(document_number=form.instance.document_number,
                          ministry=form.instance.ministry,
                          issued_date=form.instance.issued_date,
                          owner=form.instance.owner)

            with open(doc_path, 'rb') as f:
                cites.cites_file.save(
                    f'{form.instance.owner.username}_CITES_{cites.document_number}.docx', File(f))
                cites.save()

            falcon.CITES = cites
            falcon.save()

        return res

    def get_context_data(self, **kwargs):
        context = super(CITESCreate, self).get_context_data(**kwargs)
        context["form"].fields["ministry"].queryset = Office.objects.filter(
            office_type="MIN")
        context["form"].fields["falcons"].queryset = Falcon.objects.all()
        context["title"] = "Add a new CITES"
        return context


class CITESDetail(LoginRequiredMixin, UserPassesTestMixin, generic.DetailView):

    model = CITES

    def test_func(self):
        CITES = self.get_object()
        return self.request.user == CITES.owner

    def get_context_data(self, **kwargs):
        context = super(CITESDetail, self).get_context_data(**kwargs)
        context["title"] = "CITES details"
        return context


class CITESUpdate(LoginRequiredMixin, UserPassesTestMixin, generic.edit.UpdateView):

    model = CITES
    form_class = CITESUpdateForm

    def test_func(self):
        CITES = self.get_object()
        return self.request.user == CITES.owner

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(CITESUpdate, self).get_context_data(**kwargs)
        context["title"] = "Update a CITES"
        return context


class CITESDelete(LoginRequiredMixin, UserPassesTestMixin, generic.DeleteView):

    model = CITES
    success_url = "/breeding/cites"

    def test_func(self):
        CITES = self.get_object()
        return self.request.user == CITES.owner

    def get_context_data(self, **kwargs):
        context = super(CITESDelete, self).get_context_data(**kwargs)
        context["title"] = "Delete a CITES"
        return context


class OfficeList(LoginRequiredMixin, generic.ListView):

    model = Office
    paginate_by = 10

    def get_queryset(self):
        return Office.objects.filter(owner=self.request.user)

    def get_context_data(self, **kwargs):
        context = super(OfficeList, self).get_context_data(**kwargs)
        context["title"] = "List of all offices"
        return context


class OfficeDetail(LoginRequiredMixin, UserPassesTestMixin, generic.DetailView):

    model = Office

    def test_func(self):
        Office = self.get_object()
        return self.request.user == Office.owner

    def get_context_data(self, **kwargs):
        context = super(OfficeDetail, self).get_context_data(**kwargs)
        context["title"] = "Office details"
        return context


class OfficeCreate(LoginRequiredMixin, generic.edit.CreateView):
    """ View for creating new Office """

    model = Office
    form_class = OfficeCreateForm

    def form_valid(self, form):
        form.instance.owner = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(OfficeCreate, self).get_context_data(**kwargs)
        context["title"] = "Add a new Office"
        return context


class OfficeUpdate(LoginRequiredMixin, UserPassesTestMixin, generic.edit.UpdateView):

    model = Office
    form_class = OfficeUpdateForm

    def test_func(self):
        Office = self.get_object()
        return self.request.user == Office.owner

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(OfficeUpdate, self).get_context_data(**kwargs)
        context["title"] = "Update a Office"
        return context


class OfficeDelete(LoginRequiredMixin, UserPassesTestMixin, generic.DeleteView):

    model = Office
    success_url = "/breeding/offices"

    def test_func(self):
        Office = self.get_object()
        return self.request.user == Office.owner

    def get_context_data(self, **kwargs):
        context = super(OfficeDelete, self).get_context_data(**kwargs)
        context["title"] = "Delete an Office"
        return context
